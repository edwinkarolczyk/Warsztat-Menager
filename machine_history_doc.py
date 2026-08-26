# version: 1.1
"""Dodatkowy zapis historii przeglądów i napraw maszyny do pliku DOCX."""

from __future__ import annotations

import datetime as dt
import json
import logging
import os
import unicodedata
from pathlib import Path
from typing import Any, Iterable, Mapping, Sequence

logger = logging.getLogger(__name__)

_DONE_STATUSES = {
    "done",
    "wykonany",
    "wykonane",
    "completed",
    "zamkniety",
    "zamknięty",
}


class ServiceHistoryDocumentError(RuntimeError):
    """Błąd obsługi zewnętrznej karty historii maszyny."""


class InvalidDocumentFormatError(ServiceHistoryDocumentError):
    """Wybrany plik nie jest dokumentem DOCX."""


class HistoryTableNotFoundError(ServiceHistoryDocumentError):
    """W dokumencie nie znaleziono tabeli historii."""


def _normalize_text(value: object) -> str:
    text = str(value or "").replace("\n", " ").replace("\xa0", " ").strip()
    text = unicodedata.normalize("NFKD", text)
    text = "".join(ch for ch in text if not unicodedata.combining(ch))
    return " ".join(text.casefold().split())


def _classify_header(value: object) -> str | None:
    text = _normalize_text(value)
    if not text:
        return None
    if "podpis" in text:
        return "signature"
    if "uwagi" in text or "opis" in text:
        return "notes"
    if "data" in text:
        return "date"
    if "typ" in text:
        return "type"
    return None


def _header_mapping_for_row(row) -> dict[str, int]:
    mapping: dict[str, int] = {}
    for idx, cell in enumerate(row.cells):
        field = _classify_header(cell.text)
        if field and field not in mapping:
            mapping[field] = idx
    return mapping


def _find_history_table(document):
    required = {"type", "date", "signature", "notes"}
    for table in document.tables:
        rows = list(table.rows)
        if not rows:
            continue

        for row_idx, row in enumerate(rows[:8]):
            mapping = _header_mapping_for_row(row)
            if required.issubset(mapping):
                return table, row_idx, mapping

        max_header_rows = min(4, len(rows))
        max_columns = max(
            (len(row.cells) for row in rows[:max_header_rows]), default=0
        )
        for end_idx in range(max_header_rows):
            mapping: dict[str, int] = {}
            for col_idx in range(max_columns):
                parts = []
                for row_idx in range(end_idx + 1):
                    row = rows[row_idx]
                    if col_idx < len(row.cells):
                        parts.append(row.cells[col_idx].text)
                field = _classify_header(" ".join(parts))
                if field and field not in mapping:
                    mapping[field] = col_idx
            if required.issubset(mapping):
                return table, end_idx, mapping

    raise HistoryTableNotFoundError(
        "Nie znaleziono tabeli z kolumnami TYP, DATA, PODPIS i UWAGI/OPIS."
    )


def _format_date(value: object) -> str:
    if isinstance(value, dt.datetime):
        return value.strftime("%d.%m.%y")
    if isinstance(value, dt.date):
        return value.strftime("%d.%m.%y")

    raw = str(value or "").strip()
    if not raw:
        return dt.date.today().strftime("%d.%m.%y")
    try:
        parsed = dt.datetime.fromisoformat(raw.replace("Z", "+00:00"))
        return parsed.strftime("%d.%m.%y")
    except Exception:
        pass
    try:
        parsed_date = dt.date.fromisoformat(raw[:10])
        return parsed_date.strftime("%d.%m.%y")
    except Exception:
        return raw


def _person_signature(value: object) -> str:
    raw = str(value or "").strip()
    if not raw:
        return ""
    words = [part for part in raw.replace("-", " ").split() if part]
    if len(words) >= 2:
        return "".join(word[0].upper() for word in words)
    return raw


def format_signatures(people: object) -> str:
    if isinstance(people, str):
        values = [part.strip() for part in people.split(",") if part.strip()]
    elif isinstance(people, Iterable):
        values = [str(item).strip() for item in people if str(item or "").strip()]
    else:
        values = [str(people).strip()] if str(people or "").strip() else []
    return ", ".join(_person_signature(value) for value in values)


def _row_is_empty(row, columns: Sequence[int]) -> bool:
    for idx in columns:
        if idx < len(row.cells) and str(row.cells[idx].text or "").strip():
            return False
    return True


def _set_cell_text_preserving_basic_format(cell, value: object) -> None:
    text = str(value or "")
    if not cell.paragraphs:
        cell.text = text
        return

    paragraph = cell.paragraphs[0]
    for current in cell.paragraphs:
        for run in current.runs:
            run.text = ""

    if paragraph.runs:
        paragraph.runs[0].text = text
    else:
        paragraph.add_run(text)


def _target_data_row(table, header_row_idx: int, columns: Sequence[int]):
    rows = list(table.rows)
    for row in rows[header_row_idx + 1 :]:
        if _row_is_empty(row, columns):
            return row
    return table.add_row()


def append_history_entry(
    path: str | os.PathLike[str],
    *,
    entry_type: str,
    performed_at: object,
    performed_by: object,
    description: object,
) -> None:
    """Dopisz jeden wpis P/N do tabeli historii w dokumencie DOCX."""

    target = Path(path)
    if target.suffix.casefold() != ".docx":
        raise InvalidDocumentFormatError("Karta historii musi być plikiem .docx.")
    if not target.is_file():
        raise ServiceHistoryDocumentError(f"Plik karty nie istnieje: {target}")

    typ = str(entry_type or "").strip().upper()
    if typ not in {"P", "N"}:
        raise ServiceHistoryDocumentError("Typ wpisu musi mieć wartość P albo N.")

    try:
        from docx import Document  # type: ignore
    except Exception as exc:
        raise ServiceHistoryDocumentError(
            "Brak biblioteki python-docx wymaganej do zapisu karty historii."
        ) from exc

    try:
        document = Document(str(target))
        table, header_idx, mapping = _find_history_table(document)
        columns = [mapping[key] for key in ("type", "date", "signature", "notes")]
        row = _target_data_row(table, header_idx, columns)

        values: Mapping[str, str] = {
            "type": typ,
            "date": _format_date(performed_at),
            "signature": format_signatures(performed_by),
            "notes": str(description or "").strip(),
        }
        for field, value in values.items():
            col_idx = mapping[field]
            if col_idx >= len(row.cells):
                raise ServiceHistoryDocumentError(
                    "Tabela historii ma niezgodną liczbę kolumn."
                )
            _set_cell_text_preserving_basic_format(row.cells[col_idx], value)

        temp_path = target.with_name(f".{target.stem}.wm_tmp{target.suffix}")
        try:
            document.save(str(temp_path))
            os.replace(temp_path, target)
        finally:
            if temp_path.exists():
                try:
                    temp_path.unlink()
                except OSError:
                    pass
    except ServiceHistoryDocumentError:
        raise
    except PermissionError as exc:
        raise ServiceHistoryDocumentError(
            "Nie można zapisać karty. Zamknij dokument w Wordzie i spróbuj ponownie."
        ) from exc
    except Exception as exc:
        raise ServiceHistoryDocumentError(
            f"Nie udało się zapisać karty historii: {exc}"
        ) from exc


def _machine_id(machine: Mapping[str, Any]) -> str:
    return str(
        machine.get("id")
        or machine.get("nr_ewid")
        or machine.get("nr")
        or ""
    ).strip()


def _status_key(value: object) -> str:
    raw = _normalize_text(value).replace("_", " ").replace("-", " ")
    if raw in {"warn", "warm", "warning", "awaria", "stop"}:
        return "warn"
    if raw in {"ok", "sprawna", "sprawny", "sprawne"}:
        return "ok"
    if raw in {"alert", "serwis", "przeglad", "serwis/przeglad"}:
        return "alert"
    return raw


def _review_key(review: Mapping[str, Any]) -> tuple[str, ...]:
    review_id = str(review.get("id") or "").strip()
    if review_id:
        return ("id", review_id)
    return (
        "fallback",
        str(review.get("planned_date") or review.get("date") or "").strip(),
        str(review.get("type") or review.get("typ") or "").strip(),
    )


def _review_done(review: Mapping[str, Any]) -> bool:
    return _normalize_text(review.get("status")) in _DONE_STATUSES


def _is_recent_timestamp(value: object, *, seconds: int = 300) -> bool:
    raw = str(value or "").strip()
    if not raw:
        return False
    try:
        parsed = dt.datetime.fromisoformat(raw.replace("Z", "+00:00"))
    except Exception:
        return False
    if parsed.tzinfo is not None:
        now = dt.datetime.now(parsed.tzinfo)
    else:
        now = dt.datetime.now()
    delta = (now - parsed).total_seconds()
    return -5 <= delta <= seconds


def _history_item_key(item: Mapping[str, Any]) -> tuple[str, ...]:
    return (
        str(item.get("status") or "").strip(),
        str(item.get("started_at") or "").strip(),
        str(item.get("ended_at") or "").strip(),
        str(item.get("closed_by") or item.get("changed_by") or "").strip(),
    )


def _read_machine_rows(path: str) -> list[dict[str, Any]]:
    try:
        with open(path, "r", encoding="utf-8-sig") as handle:
            payload = json.load(handle)
    except Exception:
        return []
    if isinstance(payload, dict) and isinstance(payload.get("maszyny"), list):
        raw_rows = payload["maszyny"]
    elif isinstance(payload, list):
        raw_rows = payload
    else:
        return []
    return [item for item in raw_rows if isinstance(item, dict)]


def _collect_history_events(
    old_rows: Sequence[Mapping[str, Any]],
    new_rows: Sequence[Mapping[str, Any]],
) -> list[dict[str, Any]]:
    old_by_id = {_machine_id(row): row for row in old_rows if _machine_id(row)}
    events: list[dict[str, Any]] = []

    for machine in new_rows:
        machine_id = _machine_id(machine)
        if not machine_id:
            continue
        old_machine = old_by_id.get(machine_id)

        old_reviews: dict[tuple[str, ...], Mapping[str, Any]] = {}
        if old_machine and isinstance(old_machine.get("reviews"), list):
            old_reviews = {
                _review_key(item): item
                for item in old_machine["reviews"]
                if isinstance(item, dict)
            }
        new_reviews = machine.get("reviews")
        if isinstance(new_reviews, list):
            for review in new_reviews:
                if not isinstance(review, dict) or not _review_done(review):
                    continue
                previous = old_reviews.get(_review_key(review))
                transitioned = previous is not None and not _review_done(previous)
                newly_completed = (
                    previous is None
                    and _is_recent_timestamp(review.get("completed_at"))
                )
                if not (transitioned or newly_completed):
                    continue
                events.append(
                    {
                        "machine": machine,
                        "type": "P",
                        "performed_at": review.get("completed_at")
                        or review.get("done_at")
                        or dt.datetime.now().isoformat(),
                        "performed_by": review.get("completed_by") or [],
                        "description": review.get("result_note")
                        or review.get("description")
                        or review.get("type")
                        or "Przegląd / serwis",
                    }
                )

        if _status_key(machine.get("status")) != "ok":
            continue

        old_history_keys: set[tuple[str, ...]] = set()
        if old_machine and isinstance(old_machine.get("status_history"), list):
            old_history_keys = {
                _history_item_key(item)
                for item in old_machine["status_history"]
                if isinstance(item, dict)
            }
        new_history = machine.get("status_history")
        if not isinstance(new_history, list):
            continue
        for item in new_history:
            if not isinstance(item, dict):
                continue
            if _status_key(item.get("status")) != "warn" or not item.get("ended_at"):
                continue
            item_key = _history_item_key(item)
            is_new = item_key not in old_history_keys
            if old_machine is None and not _is_recent_timestamp(item.get("ended_at")):
                is_new = False
            if not is_new:
                continue
            events.append(
                {
                    "machine": machine,
                    "type": "N",
                    "performed_at": item.get("ended_at"),
                    "performed_by": [
                        str(item.get("closed_by") or item.get("changed_by") or "").strip()
                    ],
                    "description": item.get("close_note")
                    or item.get("note")
                    or "Naprawa / przywrócenie sprawności",
                }
            )

    return events


def _dialog_parent(gui_module):
    tk_module = getattr(gui_module, "tk", None)
    base = getattr(tk_module, "_wm_base_tk", tk_module)
    return getattr(base, "_default_root", None)


def _resolve_docx_absolute(gui_module, stored_path: str) -> str:
    resolver = getattr(gui_module, "_resolve_card_absolute", None)
    if callable(resolver):
        try:
            return str(resolver(stored_path) or "")
        except Exception:
            pass
    return os.path.abspath(os.path.expanduser(stored_path))


def _store_docx_path(gui_module, selected_path: str) -> str:
    resolver = getattr(gui_module, "_resolve_card_storage", None)
    if callable(resolver):
        try:
            return str(resolver(selected_path) or selected_path)
        except Exception:
            pass
    return os.path.normpath(selected_path)


def _warn_gui(gui_module, title: str, text: str, *, parent=None) -> None:
    messagebox = getattr(gui_module, "messagebox", None)
    if messagebox is None:
        logger.warning("%s: %s", title, text)
        return
    try:
        messagebox.showwarning(
            title,
            text,
            parent=parent or _dialog_parent(gui_module),
        )
    except Exception:
        logger.warning("%s: %s", title, text)


def _choose_docx_for_machine(
    machine: dict[str, Any], gui_module, *, parent=None
) -> str:
    filedialog = getattr(gui_module, "filedialog", None)
    if filedialog is None:
        return ""
    try:
        selected = filedialog.askopenfilename(
            parent=parent or _dialog_parent(gui_module),
            title="Wybierz kartę historii maszyny (.docx)",
            filetypes=(("Dokument Word DOCX", "*.docx"),),
        )
    except Exception:
        return ""
    if not selected:
        return ""
    if Path(selected).suffix.casefold() != ".docx":
        _warn_gui(
            gui_module,
            "Karta historii maszyny",
            "Można wybrać tylko plik w formacie .docx.",
            parent=parent,
        )
        return ""
    stored = _store_docx_path(gui_module, selected)
    machine["service_history_file"] = stored
    return stored


def _valid_assigned_docx(
    machine: dict[str, Any],
    gui_module,
    *,
    parent=None,
    prompt_if_missing: bool,
) -> tuple[str, bool]:
    stored = str(machine.get("service_history_file") or "").strip()
    newly_assigned = False
    if not stored and prompt_if_missing:
        stored = _choose_docx_for_machine(machine, gui_module, parent=parent)
        newly_assigned = bool(stored)
    if not stored:
        return "", newly_assigned
    if Path(stored).suffix.casefold() != ".docx":
        _warn_gui(
            gui_module,
            "Karta historii maszyny",
            "Przypisana karta historii nie jest plikiem .docx. "
            "Wybierz poprawny plik przy edycji maszyny.",
            parent=parent,
        )
        return "", newly_assigned
    absolute = _resolve_docx_absolute(gui_module, stored)
    if not absolute or not os.path.isfile(absolute):
        _warn_gui(
            gui_module,
            "Karta historii maszyny",
            "Historia WM została zapisana, ale przypisany plik karty nie istnieje:\n"
            f"{absolute}",
            parent=parent,
        )
        return "", newly_assigned
    return absolute, newly_assigned


def _write_events_after_machine_save(
    events: Sequence[Mapping[str, Any]],
    gui_module,
    original_write,
    path: str,
    payload: Any,
) -> None:
    path_changed = False
    for event in events:
        machine = event.get("machine")
        if not isinstance(machine, dict):
            continue
        absolute, newly_assigned = _valid_assigned_docx(
            machine,
            gui_module,
            prompt_if_missing=True,
        )
        path_changed = path_changed or newly_assigned
        if not absolute:
            continue
        try:
            append_history_entry(
                absolute,
                entry_type=str(event.get("type") or ""),
                performed_at=event.get("performed_at"),
                performed_by=event.get("performed_by") or [],
                description=event.get("description") or "",
            )
        except Exception as exc:
            logger.exception(
                "[Maszyny][DOCX_HISTORY] Nie udało się dopisać historii do %s",
                absolute,
            )
            _warn_gui(
                gui_module,
                "Karta historii maszyny",
                "Historia w WM została zapisana, ale nie udało się dopisać "
                f"wpisu do karty DOCX:\n{exc}",
            )

    if path_changed:
        try:
            original_write(path, payload)
        except Exception:
            logger.exception(
                "[Maszyny][DOCX_HISTORY] Nie udało się utrwalić ścieżki karty."
            )


def _install_machine_save_hook(gui_module) -> None:
    original = getattr(gui_module, "_safe_write_json", None)
    if not callable(original) or getattr(original, "_wm_docx_history_wrapper", False):
        return

    def _wrapped_write(path: str, data: Any) -> bool:
        rows = data.get("maszyny") if isinstance(data, dict) else None
        if not isinstance(rows, list):
            return bool(original(path, data))

        old_rows = _read_machine_rows(path)
        new_rows = [row for row in rows if isinstance(row, dict)]
        events = _collect_history_events(old_rows, new_rows)
        result = bool(original(path, data))
        if result and events:
            _write_events_after_machine_save(
                events,
                gui_module,
                original,
                path,
                data,
            )
        return result

    _wrapped_write._wm_docx_history_wrapper = True  # type: ignore[attr-defined]
    _wrapped_write._wm_docx_original = original  # type: ignore[attr-defined]
    gui_module._safe_write_json = _wrapped_write


def _decorate_machine_edit_dialog(window, gui_module) -> None:
    if getattr(window, "_wm_docx_history_decorated", False):
        return
    if window.__class__.__name__ != "MachineEditDialog":
        return
    row = getattr(window, "_row", None)
    if not isinstance(row, dict):
        return

    window._wm_docx_history_decorated = True
    ttk = getattr(gui_module, "ttk", None)
    if ttk is None:
        return

    box = ttk.LabelFrame(
        window,
        text="Dodatkowa karta historii przeglądów i napraw",
    )
    box.pack(side="bottom", fill="x", padx=12, pady=(4, 0))

    label = ttk.Label(box)
    label.pack(side="left", padx=(8, 12), pady=6)

    open_button = ttk.Button(box, text="Otwórz kartę")
    open_button.pack(side="left", pady=6)

    def _refresh() -> None:
        stored = str(row.get("service_history_file") or "").strip()
        if stored:
            label.configure(text=f"Plik: {os.path.basename(stored)}")
            open_button.state(["!disabled"])
        else:
            label.configure(text="Plik: brak")
            open_button.state(["disabled"])

    def _choose() -> None:
        stored = _choose_docx_for_machine(row, gui_module, parent=window)
        if not stored:
            return
        try:
            window._dirty = True
        except Exception:
            pass
        _refresh()

    def _open() -> None:
        stored = str(row.get("service_history_file") or "").strip()
        if not stored:
            return
        absolute = _resolve_docx_absolute(gui_module, stored)
        if not absolute or not os.path.isfile(absolute):
            _warn_gui(
                gui_module,
                "Karta historii maszyny",
                f"Plik nie istnieje:\n{absolute}",
                parent=window,
            )
            return
        opener = getattr(gui_module, "_open_external", None)
        if not callable(opener) or not opener(absolute):
            _warn_gui(
                gui_module,
                "Karta historii maszyny",
                "Nie udało się otworzyć pliku.",
                parent=window,
            )

    ttk.Button(
        box,
        text="Wybierz plik .docx...",
        command=_choose,
    ).pack(side="left", padx=(0, 6), pady=6, before=open_button)
    open_button.configure(command=_open)

    original_on_ok = getattr(window, "_on_ok", None)
    if callable(original_on_ok) and not getattr(
        original_on_ok,
        "_wm_docx_history_wrapper",
        False,
    ):

        def _on_ok_with_history(updated_row):
            stored = str(row.get("service_history_file") or "").strip()
            if isinstance(updated_row, dict) and stored:
                updated_row["service_history_file"] = stored
            return original_on_ok(updated_row)

        _on_ok_with_history._wm_docx_history_wrapper = True  # type: ignore[attr-defined]
        window._on_ok = _on_ok_with_history

    _refresh()


def _install_machine_edit_dialog_hook(gui_module) -> None:
    tk_module = getattr(gui_module, "tk", None)
    if tk_module is None or getattr(tk_module, "_wm_docx_history_proxy", False):
        return

    real_toplevel = getattr(tk_module, "Toplevel", None)
    if real_toplevel is None:
        return

    class _HistoryAwareToplevel(real_toplevel):
        def __init__(self, *args, **kwargs):
            super().__init__(*args, **kwargs)
            try:
                self.after_idle(
                    lambda: _decorate_machine_edit_dialog(self, gui_module)
                )
            except Exception:
                pass

    class _TkProxy:
        _wm_docx_history_proxy = True
        _wm_base_tk = tk_module
        Toplevel = _HistoryAwareToplevel

        def __getattr__(self, name: str):
            return getattr(tk_module, name)

    gui_module.tk = _TkProxy()


def install_gui_integration(gui_module) -> bool:
    """Podłącz dodatkową kartę DOCX do aktualnego modułu ``gui_maszyny``."""

    if gui_module is None:
        return False
    if getattr(gui_module, "_wm_docx_history_installed", False):
        return True

    required = (
        "_safe_write_json",
        "_resolve_card_storage",
        "_resolve_card_absolute",
        "_open_external",
        "filedialog",
        "messagebox",
        "ttk",
        "tk",
    )
    if any(not hasattr(gui_module, name) for name in required):
        return False

    _install_machine_save_hook(gui_module)
    _install_machine_edit_dialog_hook(gui_module)
    gui_module._wm_docx_history_installed = True
    return True


__all__ = [
    "ServiceHistoryDocumentError",
    "InvalidDocumentFormatError",
    "HistoryTableNotFoundError",
    "append_history_entry",
    "format_signatures",
    "install_gui_integration",
]
